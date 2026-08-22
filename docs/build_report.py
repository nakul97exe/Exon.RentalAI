"""Build docs/Exon_Rental_Capstone_Report.docx — the project documentation.

Run from the repo root with the venv active:

    python docs/build_report.py

Screenshots: drop PNG/JPG files into docs/images/ using the names listed in
SHOTS below. Any that are missing render as a labelled placeholder box, so the
document builds either way and you can add images incrementally.
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS = Path(__file__).resolve().parent
IMAGES = DOCS / "images"
OUT = DOCS / "Exon_Rental_Capstone_Report.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x59, 0x59, 0x59)
CODE_BG = "F2F4F7"
SHOT_BG = "EDEFF2"

# ---------------------------------------------------------------- primitives


def shade(cell, hex_fill):
    """Fill a table cell with a solid colour."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ACCENT
    return p


def para(doc, text, italic=False, size=10.5, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    if italic:
        run.font.color.rgb = MUTED
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullets(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        # "**bold**: rest" renders the label in bold.
        if item.startswith("**") and "**" in item[2:]:
            label, rest = item[2:].split("**", 1)
            p.add_run(label).bold = True
            p.add_run(rest).font.size = Pt(10.5)
            p.runs[0].font.size = Pt(10.5)
        else:
            p.add_run(item).font.size = Pt(10.5)


def code(doc, text, caption=None):
    """A monospaced, shaded block. One cell so it never splits oddly."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade(cell, CODE_BG)

    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for line in text.strip("\n").split("\n"):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
    if caption:
        para(doc, caption, italic=True, size=9, space_after=10)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, name in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade(cell, "DCE6F1")
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(name)
        run.bold = True
        run.font.size = Pt(9.5)

    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].paragraphs[0].text = ""
            run = cells[i].paragraphs[0].add_run(str(value))
            run.font.size = Pt(9.5)

    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def screenshot(doc, filename, caption):
    """Embed docs/images/<filename>, or a placeholder box if it isn't there."""
    path = IMAGES / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = t.rows[0].cells[0]
        shade(cell, SHOT_BG)
        cell.paragraphs[0].text = ""
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(
            f"\n[ screenshot pending — save as docs/images/{filename} ]\n"
        )
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = MUTED

    p = para(
        doc,
        caption,
        italic=True,
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )
    return p


SHOTS = [
    ("01-main-shell.png", "Figure 1 — Application shell: icon rail, Esri map, chat panel."),
    ("02-documents-panel.png", "Figure 2 — Documents drawer: ordinance upload and ingestion result."),
    ("03-add-data-panel.png", "Figure 3 — Add data drawer: parcel shapefile (.zip) upload."),
    ("04-parcel-selected.png", "Figure 4 — A selected parcel and its attribute popup."),
    ("05-answer-with-sources.png", "Figure 5 — A grounded answer with source chips and the verification badge."),
    ("06-reasoning-trace.png", "Figure 6 — The expanded reasoning trace: plan, queries, attempts."),
]

# ---------------------------------------------------------------- document

doc = Document()

base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(10.5)

for s in doc.sections:
    s.top_margin = Inches(0.9)
    s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)

# ---- cover
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Exon Rental")
run.bold = True
run.font.size = Pt(34)
run.font.color.rgb = ACCENT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Agentic RAG Assistant for Rental-Housing Compliance")
run.font.size = Pt(15)
run.font.color.rgb = MUTED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Project Documentation")
run.font.size = Pt(13)
run.bold = True

doc.add_paragraph()
para(
    doc,
    "Municipal ordinances joined to GIS parcel data, so the answer depends on the "
    "specific property. Demo city: Palo Alto, California.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
)

doc.add_paragraph()
table(
    doc,
    ["Field", "Detail"],
    [
        ("Project", "Exon Rental — parcel-aware housing compliance assistant"),
        ("Document", "Project documentation (architecture, flows, code, use cases)"),
        ("Version", "1.0"),
        ("Status", "Code complete and deployed"),
        ("Backend", "FastAPI on Azure App Service (Linux, Python 3.11)"),
        ("Frontend", "React 19 + Vite on Azure Static Web Apps"),
        ("Repository", "paloAltoRentalGIS (branch: main)"),
        ("Live demo", "https://red-beach-0bbb9cb1e.7.azurestaticapps.net/"),
        ("Sample files", "docs/files/ in the repository — the ordinance PDF and parcel shapefile used throughout this document"),
    ],
    widths=[1.5, 4.7],
)

doc.add_page_break()

# ---- index
h(doc, "Index", 1)
para(
    doc,
    "Section numbers below match the headings in the body of this document.",
    italic=True,
)

INDEX = [
    ("1", "Overview", "Purpose, the problem, what the system does, key features"),
    ("2", "Technology Stack", "Backend, frontend, infrastructure and deliberate non-choices"),
    ("3", "System Architecture", "Components, responsibilities, the agent pipeline"),
    ("4", "Application Flows", "Document upload, parcel-layer upload, question answering"),
    ("5", "Code Structure", "Repository tree and module-by-module reference"),
    ("6", "Code Walkthrough", "Annotated excerpts of the load-bearing code"),
    ("7", "User Interface", "Screens and what each one does"),
    ("8", "Use Cases", "Two end-to-end scenarios with expected results"),
    ("9", "API Reference", "Endpoints, request and response shapes"),
    ("10", "Deployment", "CI/CD, hosting, startup command, app settings"),
    ("11", "Running Locally", "Setup commands and environment variables"),
]
table(doc, ["#", "Section", "Contents"], INDEX, widths=[0.4, 1.9, 3.9])

doc.add_page_break()

# ---- 1. overview
h(doc, "1. Overview", 1)

h(doc, "1.1 Purpose", 2)
para(
    doc,
    "Exon Rental answers plain-English questions about a city's rental-housing "
    "ordinance, and grounds every answer in two sources at once: the ordinance text "
    "itself and the attributes of a specific parcel selected on a map. It was built "
    "as a postgraduate AI/ML capstone project.",
)

h(doc, "1.2 The problem", 2)
para(
    doc,
    "A tenant facing eviction wants to know what they are owed. The answer lives in "
    "two places that never talk to each other:",
)
bullets(
    doc,
    [
        "**The ordinance** — a PDF of municipal code, written in legal prose.",
        "**The property record** — a GIS parcel dataset holding unit counts, zoning and land use.",
    ],
)
para(
    doc,
    "Neither alone answers the question. The ordinance says buildings with 10 or more "
    "units pay $13,000 for a two-bedroom; it does not know how many units this "
    "building has. The GIS layer knows the unit count but nothing about the law. A "
    "generic document chatbot reads the PDF and stops there — it will quote a figure "
    "for a rule that does not apply to the building in question. This project joins "
    "the two.",
)

h(doc, "1.3 What the system does", 2)
table(
    doc,
    ["Step", "Action", "Result"],
    [
        ("1", "Upload an ordinance (PDF / TXT / CSV / Excel)", "Chunked section-by-section, embedded, stored in a vector database"),
        ("2", "Upload a parcel shapefile (.zip)", "Parsed in the browser and drawn on an Esri map"),
        ("3", "Click a parcel", "Its attributes become the context for the next question"),
        ("4", "Ask a question in plain English", "Four agents plan, retrieve, reason, generate and verify the answer"),
    ],
    widths=[0.45, 2.35, 3.4],
)
para(doc, "Every answer carries three things beyond the text itself:")
bullets(
    doc,
    [
        "**Source chips** — the exact ordinance sections used, with relevance scores.",
        "**A verification badge** — whether every claim traces back to a retrieved source.",
        "**A reasoning trace** — the plan, the tools chosen and the queries attempted.",
    ],
)

h(doc, "1.4 Key features", 2)
table(
    doc,
    ["Feature", "Description"],
    [
        ("Parcel-aware answers", "The ordinance section applied is chosen from the selected parcel's real unit count."),
        ("Section-aware chunking", "A legal section is kept whole, so a rule never arrives without its exemptions."),
        ("Self-correcting retrieval", "The reasoner judges its own retrieval and searches again with a refined query, up to two retries."),
        ("Independent validation", "A separate agent checks the answer against the sources and flags unsupported claims."),
        ("Explainability", "Plan, queries, attempts, sources and scores are all returned to the UI."),
        ("City-agnostic ingestion", "The heading regex and parcel ID detection are generic, not Palo Alto specific."),
        ("Local embeddings", "Embeddings run on the server with sentence-transformers; the only paid dependency is the chat model."),
        ("No server-side parcel store", "Parcel attributes travel with each request, so concurrent users cannot see each other's selection."),
    ],
    widths=[1.7, 4.5],
)

para(
    doc,
    "In one line: click a 30-unit building and the answer is $13,000; click the "
    "single-family house next door and it is one month's rent. Same question, "
    "different answer, both correct, both cited.",
    italic=True,
)

doc.add_page_break()

# ---- 2. tech stack
h(doc, "2. Technology Stack", 1)

h(doc, "2.1 Backend", 2)
table(
    doc,
    ["Layer", "Choice", "Why"],
    [
        ("API", "FastAPI", "Async-capable, auto-generated /docs, native file uploads"),
        ("Document parsing", "PyPDFLoader, TextLoader, pandas", "Covers PDF, TXT, CSV and Excel"),
        ("Chunking", "Custom section splitter + RecursiveCharacterTextSplitter", "A legal section is a self-contained rule"),
        ("Embeddings", "sentence-transformers/all-MiniLM-L6-v2 (local)", "384-dim, ~90 MB, free, no API dependency"),
        ("Vector store", "Chroma (embedded, persisted to disk)", "No server to run, no external service"),
        ("LLM", "OpenRouter → anthropic/claude-sonnet-4.5", "One key, many models, easy to swap"),
        ("Agents", "Plain Python functions", "Every reasoning step stays visible and explainable"),
    ],
    widths=[1.2, 2.2, 2.8],
)

h(doc, "2.2 Frontend", 2)
table(
    doc,
    ["Layer", "Choice"],
    [
        ("Framework", "React 19 + Vite"),
        ("Mapping", "Esri ArcGIS JS SDK 4.31 (@arcgis/core)"),
        ("Shapefile parsing", "shpjs — runs entirely in the browser"),
        ("Styling", "Plain CSS with custom properties (dark theme)"),
    ],
    widths=[1.7, 4.5],
)

h(doc, "2.3 Infrastructure", 2)
table(
    doc,
    ["Layer", "Choice"],
    [
        ("Backend host", "Azure App Service (Linux, Python 3.11)"),
        ("Frontend host", "Azure Static Web Apps"),
        ("CI/CD", "GitHub Actions — separate ci.yml and deploy.yml"),
        ("Auth to Azure", "Publish profile (backend), deployment token (frontend)"),
    ],
    widths=[1.7, 4.5],
)

h(doc, "2.4 Deliberate non-choices", 2)
bullets(
    doc,
    [
        "**No agent framework.** LangChain supplies loaders, splitters, embeddings and the Chroma wrapper. The agent loop itself is hand-written so every reasoning step is inspectable.",
        "**No external embedding API.** Embeddings run locally on the server.",
        "**No relational database.** Chroma persists to disk; parcel attributes are never stored server-side.",
    ],
)

doc.add_page_break()

# ---- 3. architecture
h(doc, "3. System Architecture", 1)

h(doc, "3.1 Components", 2)
code(
    doc,
    """
 BROWSER                                  AZURE APP SERVICE
 ┌───────────────────────────┐            ┌──────────────────────────────────┐
 │  React 19 + Vite          │            │  FastAPI                         │
 │                           │            │                                  │
 │  Esri MapView             │  POST      │  /api/upload_document            │
 │   └ click → parcel attrs  │ ─────────► │    loader → chunker → Chroma     │
 │                           │            │                                  │
 │  shpjs (.zip → GeoJSON)   │            │  /api/query                      │
 │                           │  POST      │    orchestrator                  │
 │  ChatPanel                │ ─────────► │      planner                     │
 │   ├ answer text           │ ◄───────── │      retriever-reasoner ──┐      │
 │   ├ source chips + scores │   JSON     │      generator            │      │
 │   ├ validation badge      │            │      validator            │      │
 │   └ reasoning trace       │            │                           ▼      │
 └───────────────────────────┘            │  Chroma (on disk)   OpenRouter   │
                                          │  all-MiniLM-L6-v2   claude-4.5  │
                                          └──────────────────────────────────┘
""",
    "Figure A — Component view. The shapefile never leaves the browser; the ordinance never leaves the server.",
)

h(doc, "3.2 Layer responsibilities", 2)
table(
    doc,
    ["Layer", "Responsibility"],
    [
        ("Presentation", "Map, drawers, chat. Owns parcel selection and renders explainability output."),
        ("API routers", "Validate input, map exceptions to HTTP status codes. No business logic."),
        ("Ingestion", "File → documents → section-aware chunks. Deterministic; no LLM involved."),
        ("Vector store", "Embed, persist, retrieve. Stitches split sections back together on read."),
        ("Agents", "Plan, retrieve-and-judge, generate, validate. All LLM reasoning lives here."),
        ("Parcel data", "Clean the attributes that arrived with the request. Stateless by design."),
        ("LLM client", "One thin wrapper over OpenRouter's chat-completions endpoint."),
    ],
    widths=[1.4, 4.8],
)

h(doc, "3.3 The agent pipeline", 2)
para(
    doc,
    "Four agents run in a fixed sequence per question. Each receives only what it "
    "needs, which is what keeps the token cost and the failure modes bounded.",
)
table(
    doc,
    ["Agent", "Input", "Output", "Failure mode"],
    [
        ("Planner", "Question, has_apn, parcel field names", "needs_documents, needs_parcel, rewritten search_query", "Malformed JSON → falls back to using both tools"),
        ("Retriever-Reasoner", "Question, plan, parcel attributes", "Documents, parcel, sufficiency verdict, query history", "Judgment call fails open — reports sufficient and stops"),
        ("Generator", "Question, full section text, parcel attributes", "Grounded answer with citations", "Raises; the router returns 502"),
        ("Validator", "Answer plus the same sources", "supported flag, unsupported claims", "Fails open with checked:false — the UI shows no badge"),
    ],
    widths=[1.15, 1.75, 1.7, 1.6],
)
para(
    doc,
    "Four LLM calls per question, roughly 6–10 seconds. Plain RAG would be one call "
    "and about two seconds — that is the measured cost of the self-checking "
    "behaviour, and it is deliberate.",
    italic=True,
)

doc.add_page_break()

# ---- 4. flows
h(doc, "4. Application Flows", 1)

h(doc, "4.1 Flow — document upload", 2)
code(
    doc,
    """
User picks a PDF in the Documents drawer
   │
   ▼  DocumentsPanel.jsx → api/client.js
POST /api/upload_document          (multipart/form-data, field name "file")
   │
   ▼  routers/upload_document.py
1. Sanitise filename with Path(...).name      (blocks path traversal)
2. Validate the extension BEFORE writing to disk
3. Write to data/uploaded_documents/
   │
   ▼  ingestion/document_loader.py
4. PDF → PyPDFLoader → one Document per page
   │
   ▼  ingestion/chunker.py
5. Join the pages into one string  (a section can straddle a page break)
6. Find section headings; keep the LAST match per number to drop the TOC copy
7. Cut heading-to-heading                      → 9 sections
8. Sub-split anything over 1000 characters     → 37 chunks
   │
   ▼  vectorstore/chroma_client.py
9. Build stable ids: source_file::section::part
10. Chroma embeds each chunk (384 numbers) and persists to disk
   │
   ▼
{"filename": "...", "pages": 7, "chunks": 37, "sections": [9 numbers]}
""",
    "Figure B — Ingestion. No LLM is involved: chunking a PDF has one correct answer, so there is nothing to reason about.",
)

h(doc, "4.2 Flow — parcel layer upload", 2)
para(doc, "Entirely client-side. The backend never sees the shapefile.")
code(
    doc,
    """
User picks parcels.zip in the Add data drawer
   │
   ▼  lib/shapefile.js
1. shpjs unzips, reads .shp/.dbf/.prj and reprojects to WGS84
2. Filter to supported geometry types
3. Read the geometry type off the first feature
   (the shapefile format allows only one type per file)
4. Wrap the features in a Blob URL → new GeoJSONLayer
5. Attach a renderer and a popup template with explicit fieldInfos
   │
   ▼  AddDataPanel.jsx
6. view.map.add(layer); await layer.when()
7. Zoom to layer.fullExtent
8. Report the feature count up to App → TopBar
""",
    "Figure C — Client-side parsing. The map needs the geometry anyway, so parsing in the browser avoids uploading 20 MB and re-serving it.",
)
para(
    doc,
    "The trade-off is that the backend has no parcel database — which is exactly why "
    "attributes travel with each query.",
    italic=True,
)

h(doc, "4.3 Flow — asking a question", 2)
code(
    doc,
    """
User clicks parcel 127-53-008, types a question
   │
   ▼  MapView.jsx
view.on("click") → hitTest → { apn: "127-53-008", attributes: {...} }
   │
   ▼  App.jsx → ChatPanel.jsx → api/client.js
POST /api/query { question, apn, parcel_attributes }
   │
   ▼  routers/query.py → agents/orchestrator.py
   │
   ├── 1. PLANNER ───────────────────────────────────────── LLM call #1
   │      in:  question, has_apn, parcel field NAMES (not values)
   │      out: { needs_documents, needs_parcel, search_query, reasoning }
   │
   ├── 2. RETRIEVER-REASONER
   │      ├─ get_parcel_attributes(apn, raw)     ← once, outside the loop
   │      └─ loop, max 3 attempts:
   │           ├─ search_sections(query)          ← local embeddings, no API
   │           ├─ filter score > MIN_SCORE (0.15)
   │           ├─ summarise findings (section labels only — cheap)
   │           ├─ check_sufficiency(...)          ── LLM call #2
   │           └─ sufficient ? break : query = better_query
   │
   ├── 3. GENERATOR ─────────────────────────────────────── LLM call #3
   │      in:  question, full section text, parcel attributes
   │      out: grounded answer with citations
   │
   └── 4. VALIDATOR ─────────────────────────────────────── LLM call #4
          in:  answer + the same sources the generator saw
          out: { supported, unsupported_claims }
   │
   ▼
{ answer, sources[], parcel, validation{}, trace{} }
   │
   ▼  ChatPanel.jsx
Answer text · source chips · verification badge · collapsible trace
""",
    "Figure D — Query time. This is where every LLM call in the system happens.",
)

doc.add_page_break()

# ---- 5. code structure
h(doc, "5. Code Structure", 1)

h(doc, "5.1 Repository tree", 2)
code(
    doc,
    """
paloAltoRentalGIS/
│
├── .github/workflows/
│   ├── ci.yml                      # verify: lint, compile, import, build
│   └── deploy.yml                  # ship: backend → App Service, frontend → SWA
│
├── backend/
│   ├── app/
│   │   ├── __init__.py             # sqlite3 shim for Azure Linux (load-bearing)
│   │   ├── main.py                 # FastAPI app, CORS, router registration
│   │   ├── config.py               # env vars, paths, model names
│   │   ├── rag.py                  # plain-RAG baseline + shared prompt helpers
│   │   │
│   │   ├── routers/
│   │   │   ├── upload_document.py  # POST /api/upload_document
│   │   │   └── query.py            # POST /api/query
│   │   │
│   │   ├── ingestion/
│   │   │   ├── document_loader.py  # any file type → LangChain Documents
│   │   │   └── chunker.py          # section-aware chunking + size fallback
│   │   │
│   │   ├── vectorstore/
│   │   │   ├── embeddings.py       # cached sentence-transformers model
│   │   │   └── chroma_client.py    # add, search, stitch split sections
│   │   │
│   │   ├── parcel_data/
│   │   │   └── parcel_store.py     # cleans parcel attributes from the request
│   │   │
│   │   ├── agents/
│   │   │   ├── planner_agent.py            # which tools? rewrite the query
│   │   │   ├── retriever_reasoner_agent.py # call tools, judge, retry
│   │   │   ├── generator_agent.py          # write the grounded answer
│   │   │   ├── validator_agent.py          # verify claims against sources
│   │   │   ├── orchestrator.py             # sequence the four
│   │   │   └── json_utils.py               # shared LLM-JSON parser
│   │   │
│   │   └── llm/
│   │       └── openrouter_client.py        # requests.post to OpenRouter
│   │
│   ├── data/                       # gitignored — regenerated by uploads
│   │   ├── chroma_store/
│   │   └── uploaded_documents/
│   ├── requirements.txt
│   └── .env.example
│
├── frontend/
│   └── src/
│       ├── main.jsx                # entry; Esri assetsPath + dark theme
│       ├── App.jsx                 # app shell, shared state
│       ├── index.css               # dark theme, layout, Esri overrides
│       │
│       ├── api/client.js           # fetch wrappers for both endpoints
│       ├── lib/shapefile.js        # zip → GeoJSONLayer, ID-field detection
│       │
│       └── components/
│           ├── TopBar.jsx          # title + counts
│           ├── IconRail.jsx        # 56px icon strip
│           ├── railItems.js        # single source of truth for rail labels
│           ├── Drawer.jsx          # 320px slide-out, routes to a panel
│           ├── MapView.jsx         # Esri map + click-to-select
│           ├── ChatPanel.jsx       # conversation, sources, badge, trace
│           ├── Footer.jsx          # legend + copyright
│           ├── Icon.jsx            # inline SVG icon set
│           └── panels/
│               ├── LayersPanel.jsx         # Esri LayerList
│               ├── AddDataPanel.jsx        # shapefile upload
│               ├── DocumentsPanel.jsx      # document upload
│               ├── FindParcelPanel.jsx     # Esri Search
│               └── BaseMapGalleryPanel.jsx
│
├── docs/
│   ├── CAPSTONE_REPORT.md          # the markdown report
│   ├── build_report.py             # builds this document
│   ├── Exon_Rental_Capstone_Report.docx
│   └── images/                     # UI screenshots embedded above
│
└── .gitignore
""",
)

h(doc, "5.2 Backend module reference", 2)
table(
    doc,
    ["Module", "What it does"],
    [
        ("app/__init__.py", "Swaps pysqlite3-binary in as sqlite3, because Azure's Linux image ships a version older than Chroma's minimum (3.35). Must stay the first thing imported; a no-op locally."),
        ("app/main.py", "Creates the FastAPI app, adds CORS from config, registers routers, exposes GET /health."),
        ("app/config.py", "One place for paths and model names, all env-overridable so Azure can set DATA_DIR and ALLOWED_ORIGINS without a code change."),
        ("app/rag.py", "SYSTEM_PROMPT (the five grounding rules), MIN_SCORE = 0.15, build_context() and answer_question() — the plain-RAG baseline kept for comparison."),
        ("ingestion/document_loader.py", "Dispatches on extension: PDF → PyPDFLoader, TXT/MD → TextLoader, CSV/Excel → pandas flattened to 'col: value' lines so column names stay attached."),
        ("ingestion/chunker.py", "Section split, then size enforcement. Falls back to plain size-based chunking when a document has no recognisable numbering."),
        ("vectorstore/embeddings.py", "get_embeddings() behind lru_cache so the model loads once. normalize_embeddings=True so cosine similarity is not skewed by chunk length."),
        ("vectorstore/chroma_client.py", "Stable chunk ids, raw search, get_full_section, and search_sections — the over-fetch-dedup-stitch retrieval."),
        ("parcel_data/parcel_store.py", "Cleans what the frontend sent: keeps useful fields in a stable order, drops geometry and empty values, falls back to other short fields for unfamiliar schemas."),
        ("llm/openrouter_client.py", "chat(messages, model, temperature, max_tokens) over requests.post. temperature=0 by default — grounded legal answers should be repeatable."),
        ("routers/upload_document.py", "Sanitises the filename, validates the extension before writing, runs load → chunk → store, deletes the file if ingestion fails. A sync def on purpose: the work is CPU-bound, so FastAPI runs it in a threadpool."),
        ("routers/query.py", "Pydantic body of question, apn, parcel_attributes. Maps PlanAgentError → 400, LLMError → 502, anything else → 500."),
    ],
    widths=[1.7, 4.5],
)

h(doc, "5.3 Frontend module reference", 2)
table(
    doc,
    ["Module", "What it does"],
    [
        ("main.jsx", "Sets esriConfig.assetsPath to Esri's CDN. Without it, icons and fonts 404 — Vite cannot see runtime-constructed asset URLs, so it never bundles them."),
        ("App.jsx", "Owns three pieces of state: activePanel, view (the Esri MapView, lifted so drawer widgets can use it) and selectedParcel."),
        ("api/client.js", "fetch wrappers for both endpoints, turning FastAPI's detail field into a thrown Error."),
        ("lib/shapefile.js", "layersFromShapefileZip() unzips with shpjs, builds a GeoJSONLayer from a Blob URL and applies renderers. detectIdField() finds the parcel identifier column — APN, PARCELID, PIN, AIN."),
        ("components/MapView.jsx", "Creates the map in a useEffect with [] and returns view.destroy() for cleanup, because StrictMode double-mounts. Callbacks live in refs so a changing identity never tears down the map."),
        ("components/ChatPanel.jsx", "Conversation state, source chips, validation badge and the collapsible reasoning trace."),
        ("Esri widget panels", "LayersPanel, BaseMapGalleryPanel, FindParcelPanel and Footer mount real Esri widgets via the container property, each into a throwaway inner div — Esri's destroy() removes its container element from the DOM."),
    ],
    widths=[1.7, 4.5],
)

doc.add_page_break()

# ---- 6. code walkthrough
h(doc, "6. Code Walkthrough", 1)
para(
    doc,
    "Five excerpts, taken verbatim from the deployed code, covering the parts where "
    "the design decisions actually live.",
)

h(doc, "6.1 Section-aware chunking — ingestion/chunker.py", 2)
para(
    doc,
    "Fixed-size chunking cuts a legal rule away from its exemptions. This splits on "
    "section headings instead, and only falls back to size-based splitting when the "
    "document has no numbering the regex recognises.",
)
code(
    doc,
    '''
# Matches a heading like "9.68.060   Relocation assistance..."
# Generic (\\d+\\.\\d+\\.\\d+) so it works for any city's code.
# The 2+ spaces requirement stops mid-sentence cross-references such as
# "see Section 9.68.050 for..." from being read as headings.
SECTION_RE = re.compile(
    r"(?m)^\\s*"
    # Optional prefix used by some codes: "§ 8.52.030", "Sec. 12-45"
    r"(?:§\\s*|Sec\\.?\\s+|Section\\s+)?"
    # 2 to 4 parts, dot- or hyphen-separated: 9.68.060, 12-45, 37.02.010
    r"(?P<num>\\d+(?:[.-]\\d+){1,3})"
    r"\\s{2,}(?P<title>[A-Z][^\\n]*)"
)

def chunk_documents(docs: list[Document]) -> list[Document]:
    """Section-aware chunking with a size-based fallback."""
    full_text = "\\n".join(d.page_content for d in docs)
    source = docs[0].metadata.get("source_file", "unknown")

    sections = _split_sections(full_text, source)
    if not sections:
        # No section numbering — plain recursive splitting.
        return _splitter.split_documents(docs)

    return _enforce_size(sections)

def _split_sections(text: str, source: str) -> list[Document]:
    matches = list(SECTION_RE.finditer(text))
    if not matches:
        return []

    # Each section number appears twice: once in the table of contents, once as
    # the real heading. Keeping the LAST match per number drops the TOC copy.
    headings = sorted(
        {m.group("num"): m for m in matches}.values(), key=lambda m: m.start()
    )

    out = []
    for i, match in enumerate(headings):
        end = headings[i + 1].start() if i + 1 < len(headings) else len(text)
        body = _tidy(text[match.start() : end])

        if len(body) < 200:  # leftover TOC fragment — title with no rule text
            continue

        out.append(Document(page_content=body, metadata={
            "source_file": source,
            "section": match.group("num"),
            "title": match.group("title").strip().rstrip("."),
        }))
    return out
''',
    "Listing 1 — Section splitting. Sections over 1000 characters are then sub-split, because all-MiniLM-L6-v2 silently truncates past its 256-token window.",
)

h(doc, "6.2 Section-stitched retrieval — vectorstore/chroma_client.py", 2)
para(
    doc,
    "Section 9.68.060 splits into seven chunks. A naive search can return the dollar "
    "table without the exemption stating that the section only applies to buildings "
    "of ten units or more. Over-fetching, de-duplicating by section and re-joining "
    "the parts guarantees the rule and its exceptions arrive together.",
)
code(
    doc,
    '''
def search_sections(query: str, k: int = 4, source_file: str | None = None):
    """Search, then return whole sections instead of fragments."""
    results = []
    seen = set()

    for doc, score in search(query, k=k * 4, source_file=source_file):
        meta = doc.metadata
        key = (meta.get("source_file"), meta.get("section"))

        if key in seen:
            continue
        seen.add(key)

        if meta.get("parts", 1) > 1 and all(key):
            results.append((get_full_section(*key) or doc, score))
        else:
            results.append((doc, score))

        if len(results) >= 4:
            break

    return results


def _chunk_id(doc: Document) -> str:
    """Stable id built from the chunk's own metadata.

    Re-uploading the same document overwrites these rows instead of storing a
    second copy of every chunk.
    """
    m = doc.metadata
    return f"{m.get('source_file')}::{m.get('section', '0')}::{m.get('part', 1)}"
''',
    "Listing 2 — Retrieval that returns whole sections, plus the stable id that makes re-uploads idempotent.",
)

h(doc, "6.3 The self-correcting loop — agents/retriever_reasoner_agent.py", 2)
para(
    doc,
    "This is the part that makes the system agentic rather than plain RAG: it judges "
    "its own retrieval and searches again with a refined query. The sufficiency check "
    "sees section labels only, never full text, because it runs up to three times.",
)
code(
    doc,
    '''
    # The planner rewrites the question into ordinance wording; fall back to the
    # raw question if it didn't.
    query = plan.get("search_query") or question

    # Once, outside the loop: a dict lookup can't return something different on
    # a second try.
    parcel = None
    if plan.get("needs_parcel"):
        parcel = get_parcel_attributes(apn, parcel_attributes)

    for attempt in range(MAX_RETRIES + 1):
        attempts = attempt + 1
        # Recorded before use, so a crash still leaves a trace of what was tried.
        queries.append(query)

        if plan.get("needs_documents"):
            result = search_sections(query)
            documents = [(d, s) for d, s in result if s > MIN_SCORE]

        findings = _summarize_findings(documents, parcel)
        verdict = check_sufficiency(question, findings)
        sufficient = verdict["sufficient"]

        if sufficient:
            break

        notes.append(verdict["missing"])

        better = verdict["better_query"]
        # No new idea, or the same query again — stop rather than burn attempts.
        if not better or better in queries:
            break

        query = better
''',
    "Listing 3 — The retry loop. The sufficiency check fails open: a broken LLM call reports 'sufficient' so the loop stops rather than burning all three attempts.",
)

h(doc, "6.4 The orchestrator — agents/orchestrator.py", 2)
para(
    doc,
    "The orchestrator holds no logic of its own beyond sequencing. It owns the "
    "request data and hands each agent only what that agent needs.",
)
code(
    doc,
    '''
def answer(question: str, apn: str | None = None,
           parcel_attributes: dict | None = None) -> dict:
    """Plan, retrieve, reason, generate, validate."""
    # 1. Planner — gets the parcel's field NAMES but not their values.
    the_plan = plan(
        question,
        has_apn=bool(apn),
        parcel_fields=list(parcel_attributes.keys()) if parcel_attributes else None,
    )

    # 2. Retriever-reasoner — calls the tools, judges the results, retries search.
    found = retrieve_and_reason(question, the_plan, apn, parcel_attributes)

    # 3. Generator — writes the answer from what was gathered.
    result = generate_answer(question, found["documents"], found["parcel"])

    # 4. Validator — checks the answer against the same sources the generator saw.
    validation = validate(
        question, result["answer"], found["documents"], found["parcel"]
    )

    return {
        "answer": result["answer"],
        "sources": result["sources"],
        "parcel": found["parcel"],
        "validation": validation,
        "trace": {
            "plan": the_plan,
            "sufficient": found["sufficient"],
            "attempts": found["attempts"],
            "queries": found["queries"],
            "notes": found["notes"],
        },
    }
''',
    "Listing 4 — The four-agent sequence and the response envelope the UI consumes.",
)

h(doc, "6.5 The frontend API client — api/client.js", 2)
code(
    doc,
    '''
/** Parse a JSON response, turning FastAPI's `detail` into a thrown Error. */
async function unwrap(res) {
  let data = null;
  try {
    data = await res.json();
  } catch {
    // Non-JSON body (e.g. a proxy error page) — fall through to the status text.
  }

  if (!res.ok) {
    // fetch() only rejects on network failure, so 4xx/5xx must be checked here.
    throw new Error(data?.detail || `${res.status} ${res.statusText}`);
  }
  return data;
}

/**
 * Ask the agent a question, optionally in the context of a selected parcel.
 * Attributes are sent per request rather than stored server-side.
 */
export async function askQuestion({ question, apn = null, parcelAttributes = null }) {
  const res = await fetch(`${API_BASE}/api/query`, {
    method: "POST",
    headers: { "Content-Type": "application/json" },
    body: JSON.stringify({ question, apn, parcel_attributes: parcelAttributes }),
  });

  return unwrap(res);
}
''',
    "Listing 5 — Both endpoints go through one error-unwrapping helper, so the UI has a single failure path to render.",
)

doc.add_page_break()

# ---- 7. UI
h(doc, "7. User Interface", 1)
para(
    doc,
    "The interface is a single-screen GIS shell: a 56-pixel icon rail on the left, a "
    "320-pixel slide-out drawer for the active tool, the Esri map filling the centre, "
    "and the chat panel on the right. Nothing navigates away — the map, the parcel "
    "selection and the conversation stay visible together, which is the point.",
)

table(
    doc,
    ["Screen area", "Purpose"],
    [
        ("Top bar", "Project title, loaded-document count, loaded-feature count."),
        ("Icon rail", "Switches the drawer between Layers, Add data, Documents, Find parcel and Basemap."),
        ("Documents drawer", "Uploads an ordinance and reports pages, chunks and sections found."),
        ("Add data drawer", "Uploads a parcel shapefile .zip, parsed in the browser."),
        ("Map", "Renders the parcel layer; a click selects a parcel and shows its attributes."),
        ("Chat panel", "Question box, answer text, source chips with scores, verification badge, collapsible reasoning trace."),
        ("Footer", "Legend and copyright."),
    ],
    widths=[1.5, 4.7],
)

for filename, caption in SHOTS:
    screenshot(doc, filename, caption)

doc.add_page_break()

# ---- 8. use cases
h(doc, "8. Use Cases", 1)
para(
    doc,
    "Two scenarios, deliberately chosen as a matched pair: the same question against "
    "two different parcels must produce two different, legally correct answers. That "
    "pair is the clearest demonstration of what joining the ordinance to GIS data "
    "buys you.",
)

h(doc, "8.1 UC-01 — Relocation assistance for a large multi-family building", 2)
table(
    doc,
    ["Attribute", "Detail"],
    [
        ("Actor", "Tenant, or a housing-department caseworker"),
        ("Goal", "Find out what relocation assistance is owed for a no-fault eviction"),
        ("Preconditions", "Palo Alto ordinance ingested; parcel layer loaded on the map"),
        ("Selected parcel", "APN 127-53-008 — 30 residential rental units"),
        ("Question", "\"If a tenant with a 2-bedroom unit here is evicted without fault, how much relocation assistance is owed?\""),
    ],
    widths=[1.4, 4.8],
)
para(doc, "Main flow", size=10.5)
bullets(
    doc,
    [
        "The user clicks the parcel; the map hit-test sends the APN and attributes to the chat panel.",
        "The planner classifies the question as needing both tools and rewrites the search query into ordinance wording.",
        "The retriever fetches whole sections; the reasoner judges the retrieval sufficient on the first attempt.",
        "The generator selects the section that matches the parcel's unit count and quotes the figures verbatim.",
        "The validator re-checks every claim against the same sections and returns a verified result.",
    ],
    style="List Number",
)
para(doc, "Expected answer", size=10.5)
code(
    doc,
    """
Based on the parcel data showing 30 units, Section 9.68.060 applies rather than
Section 9.68.050, because this section governs structures with 10 or more
residential rental units. For a 2-bedroom unit, relocation assistance is
$13,000. This amount is paid in two instalments: one half ($6,500) at notice,
the remainder ($6,500) when the renter vacates. If the unit is occupied by a
low-income household, a renter aged 60 or older, a disabled renter, or a minor,
an additional $3,000 applies.

Sources:  9.68.060 (0.62)   9.68.050 (0.41)
Status:   verified — every claim traced to a retrieved section
""",
    "Listing 6 — Every figure is verbatim from the ordinance; the section choice is driven by real parcel data.",
)
para(
    doc,
    "Alternate flow: with no parcel selected, the generator asks for the unit count "
    "instead of assuming one — _format_parcel(None) writes \"no parcel selected\" "
    "explicitly for exactly this reason.",
    italic=True,
)

h(doc, "8.2 UC-02 — The same question against a single-family parcel", 2)
table(
    doc,
    ["Attribute", "Detail"],
    [
        ("Actor", "Same"),
        ("Goal", "Confirm that the answer follows the property, not the phrasing"),
        ("Selected parcel", "A single-family parcel — 1 unit, below the ten-unit threshold"),
        ("Question", "Identical to UC-01, word for word"),
    ],
    widths=[1.4, 4.8],
)
para(doc, "Expected answer", size=10.5)
code(
    doc,
    """
This parcel has 1 residential rental unit, so Section 9.68.060 — which applies
to structures with 10 or more units — does not govern. Section 9.68.050 applies
instead: relocation assistance equal to one month's rent.

Sources:  9.68.050 (0.58)   9.68.060 (0.44)
Status:   verified
""",
    "Listing 7 — Same question, different parcel, different legally correct answer. Together with UC-01 this is the proof the GIS join is doing real work.",
)
para(
    doc,
    "Guardrail cases observed alongside these two: a parking-requirement question "
    "not covered by the uploaded document is refused rather than answered with an "
    "invented number; a question about another city is refused as outside the "
    "indexed material; and the vague \"what about the money?\" triggers the retry "
    "loop and then a clarifying question rather than a guess.",
)

doc.add_page_break()

# ---- 9. API
h(doc, "9. API Reference", 1)

h(doc, "9.1 POST /api/upload_document", 2)
para(doc, "multipart/form-data, field name \"file\". Accepts PDF, TXT, MD, CSV, XLSX.")
code(
    doc,
    """
200 OK
{
  "filename": "palo-alto-9-68.pdf",
  "pages":    7,
  "chunks":   37,
  "sections": ["9.68.010", "9.68.020", "9.68.030", "9.68.040",
               "9.68.050", "9.68.060", "9.68.070", "9.68.080", "9.68.090"]
}

400  unsupported extension, or a scanned PDF with no text layer
500  ingestion failed (the uploaded file is deleted)
""",
)

h(doc, "9.2 POST /api/query", 2)
code(
    doc,
    """
Request
{
  "question": "How much relocation assistance is owed for a 2-bedroom?",
  "apn": "127-53-008",
  "parcel_attributes": { "APN": "127-53-008", "UNITS": 30, "ZONING": "RM-40" }
}

200 OK
{
  "answer":  "Based on the parcel data showing 30 units, ...",
  "sources": [ { "section": "9.68.060", "title": "...", "score": 0.62 } ],
  "parcel":  { "APN": "127-53-008", "UNITS": 30, "ZONING": "RM-40" },
  "validation": { "checked": true, "supported": true, "unsupported_claims": [] },
  "trace": {
    "plan": { "needs_documents": true, "needs_parcel": true,
              "search_query": "relocation assistance eviction without fault" },
    "sufficient": true,
    "attempts": 1,
    "queries": ["relocation assistance eviction without fault"],
    "notes": []
  }
}

400  planner rejected the request      (PlanAgentError)
502  upstream LLM failure              (LLMError)
500  anything else
""",
)

h(doc, "9.3 GET /health", 2)
code(doc, """{"status": "ok"}""")

doc.add_page_break()

# ---- 10. deployment
h(doc, "10. Deployment", 1)
code(
    doc,
    """
git push to main
   │
   ├── ci.yml
   │     ├─ backend-check:   compileall → pip install → import smoke test
   │     └─ frontend-build:  npm ci → lint → build
   │
   └── deploy.yml
         ├─ deploy-backend:  zip backend/ → Azure App Service (publish profile)
         └─ deploy-frontend: npm build with VITE_API_BASE → Static Web Apps
              (needs: deploy-backend)
""",
    "Figure E — Two workflows. CI answers \"does the code work?\" and runs on every push and PR; deploy answers \"ship it\" and runs only on main.",
)

h(doc, "10.1 Startup command", 2)
code(
    doc,
    """
gunicorn -w 1 -k uvicorn.workers.UvicornWorker --timeout 600 \\
         --bind 0.0.0.0:8000 app.main:app
""",
)
table(
    doc,
    ["Flag", "Why"],
    [
        ("-w 1", "Each worker loads its own copy of the embedding model, so memory scales linearly with worker count."),
        ("-k uvicorn.workers.UvicornWorker", "FastAPI is ASGI; gunicorn speaks WSGI by default."),
        ("--timeout 600", "The PyTorch import takes about 51 seconds; the default 30-second timeout kills the worker mid-import and loops forever."),
    ],
    widths=[1.9, 4.3],
)

h(doc, "10.2 App settings", 2)
table(
    doc,
    ["Setting", "Purpose"],
    [
        ("SCM_DO_BUILD_DURING_DEPLOYMENT=1", "Azure runs pip install on the instance, so wheels match the server platform and the upload stays around 200 KB."),
        ("DATA_DIR=/home/data", "Only /home survives a restart on App Service Linux."),
        ("HF_HOME=/home/.cache/huggingface", "Caches the embedding model so it downloads once, not on every cold start."),
        ("ALLOWED_ORIGINS", "The deployed frontend's URL; a hardcoded localhost would block every real request."),
        ("WEBSITES_CONTAINER_START_TIME_LIMIT=1800", "The default 230 seconds is shorter than a cold PyTorch start."),
        ("OPENROUTER_API_KEY", "The only paid dependency in the system."),
    ],
    widths=[2.2, 4.0],
)

doc.add_page_break()

# ---- 11. running locally
h(doc, "11. Running Locally", 1)
para(
    doc,
    "The deployed application is live at https://red-beach-0bbb9cb1e.7.azurestaticapps.net/ and needs no setup — upload a "
    "document and a parcel layer and it works in the browser. The sample ordinance PDF "
    "and parcel shapefile used throughout this document are in docs/files/ in the "
    "repository. Note that the free App Service tier sleeps when idle, so the first "
    "request after a quiet period takes a minute or so to wake the backend and load the "
    "embedding model.",
)
para(doc, "To run it locally instead:")
code(
    doc,
    """
# Backend
cd backend
python -m venv ../venv && ../venv/Scripts/activate      # Windows
pip install -r requirements.txt
cp .env.example .env                                     # add your OpenRouter key
uvicorn app.main:app --reload --port 8000

# Frontend (second terminal)
cd frontend
npm install
npm run dev
""",
)
para(
    doc,
    "Open http://localhost:5173, upload an ordinance PDF, upload a parcel shapefile "
    "zip, click a parcel and ask a question.",
)

h(doc, "11.1 Environment variables", 2)
table(
    doc,
    ["Variable", "Default", "Purpose"],
    [
        ("OPENROUTER_API_KEY", "—", "Required. The chat model credential."),
        ("OPENROUTER_MODEL", "anthropic/claude-sonnet-4.5", "Swap the model without a code change."),
        ("DATA_DIR", "backend/data", "Where uploads and the Chroma store live."),
        ("ALLOWED_ORIGINS", "http://localhost:5173", "Comma-separated CORS origins."),
        ("VITE_API_BASE", "http://localhost:8000", "Frontend build-time backend URL."),
    ],
    widths=[1.6, 1.9, 2.7],
)

doc.add_paragraph()
para(
    doc,
    "Exon Rental — parcel-aware housing compliance. Answers are informational and "
    "not legal advice.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
)

# ---------------------------------------------------------------- save
IMAGES.mkdir(exist_ok=True)
doc.save(OUT)

missing = [name for name, _ in SHOTS if not (IMAGES / name).exists()]
print(f"Wrote {OUT}")
if missing:
    print(f"Screenshot placeholders ({len(missing)}) — add to {IMAGES}:")
    for name in missing:
        print(f"  - {name}")
